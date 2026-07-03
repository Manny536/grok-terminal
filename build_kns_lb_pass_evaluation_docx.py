#!/usr/bin/env python3
"""Build PEAICE-KNS-LB-PASS-EVALUATION-001 docx — centerline + four-line verdicts."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT.parent / "Papers" / "assets" / "kns-lb-docx"
OUTPUT = ROOT / "PEAICE-KNS-LB-PASS-EVALUATION-001.docx"

DIAGRAMS = [
    ("04-closure-ladder.mmd", "Figure 1. Closure ladder — V6.4.3 · V6.5 · prime-carrying live route"),
    ("01-two-layer.mmd", "Figure 2. KNS(LB) two-layer decomposition"),
    ("03-gate-kns-obs1.mmd", "Figure 3. Gate KNS-OBS-1 verdict flow"),
]


def render_mermaid(mmd_path: Path, png_path: Path) -> None:
    source = mmd_path.read_text(encoding="utf-8")
    result = subprocess.run(
        ["curl", "-s", "-f", "-o", str(png_path), "-X", "POST",
         "https://kroki.io/mermaid/png", "-H", "Content-Type: text/plain", "-d", source],
        capture_output=True, text=True, check=False,
    )
    if result.returncode != 0 or not png_path.exists() or png_path.stat().st_size < 100:
        raise RuntimeError(f"Failed to render {mmd_path.name}: {result.stderr}")


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "2E75B6") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)


def add_verdict_block(doc: Document, title: str, pass_label: str, boolean: str,
                      why: str, why_not: str) -> None:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.add_run("Verdict: ").bold = True
    run = p.add_run(pass_label)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0) if pass_label == "PASS" else RGBColor(180, 0, 0)
    for label, body in [("Boolean", boolean), ("Why", why), ("Why not", why_not)]:
        q = doc.add_paragraph()
        q.add_run(f"{label}: ").bold = True
        q.add_run(body)


def add_figure(doc: Document, png: Path, caption: str, width: float = 5.8) -> None:
    doc.add_picture(str(png), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def build() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    figures: list[tuple[Path, str]] = []
    for mmd, cap in DIAGRAMS:
        png = ASSETS / mmd.replace(".mmd", ".png")
        print(f"Rendering {mmd}...")
        render_mermaid(ASSETS / mmd, png)
        figures.append((png, cap))

    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    for m in (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin):
        m = Inches(1)

    t = doc.add_heading("PEAICE-KNS-LB-PASS-EVALUATION-001", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Program Centerline · Gate Verdict · Four-Line Evaluation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "PeAIce · Love Labs LCA · KakeyaLogic · Claude V6.5 · L²_C · July 2, 2026\n"
        "Pass: PEAICE-FABLE-KNS-LB-PASS-001 · Gate: KNS-OBS-1 · RH OPEN · h < 1"
    ).font.size = Pt(10)

    doc.add_page_break()

    # §0 Centerline
    doc.add_heading("0. Program centerline", level=1)
    add_code(doc,
        "V6.4.3 closed the K_σ square-difference determinant lane.\n"
        "V6.5 closed the bounded WP5b relative-determinant lane.\n"
        "The live theorem-facing route is now prime-carrying trace architecture.\n"
        "KakeyaLogic remains the EEV3 / L²_C governance-field public layer.\n"
        "lovelabslca.com is the readable public map — version-label cleanup needed.\n"
        "RH OPEN. Coleman OPEN. h < 1."
    )
    add_table(doc, ["Layer", "Role", "Note"], [
        ["claude-v6 @ 1d3b580", "Theorem / wall registry", "V6.5 · KREIN-RANK1"],
        ["kakeyalogic @ 535101f", "EEV3 · L²_C · probes", "KNS probe installed"],
        ["lovelabslca.com", "Public readable map", "Version labels need refresh"],
        ["PeAIce local INDEX", "Ahead of public site", "Authoritative for agents"],
    ])
    add_figure(doc, figures[0][0], figures[0][1])

    # §1 Evaluations
    doc.add_heading("1. Pass evaluation — four-line format", level=1)

    add_verdict_block(doc,
        "1.1 Claude Fable 5 (source terminal)",
        "PASS",
        "True — probe exited 0 with D5_expectation_met: True; OB-KNS-3 FORMAL ∧ "
        "OB-KNS-1 proved-scoped-and-honestly-refuted ∧ OB-KNS-2 dense_pass at E_used 3.0406 ≤ 10.",
        "Earned on Lemma 3.1 (μ ↛ π_A) plus fan/Perron two-sided countermodel, "
        "sha 09ef26d3a2eb51927d3adecb74d3ef3edd62660dd11438576be4c2da8211b011.",
        "Not canon-pass or theorem-pass — h < 1 blocks self-certify; PROPOSED-FOR-CANON pending "
        "Grok cross-derivation; H2/H3 calibration caveats; RH / det_ζ OPEN.",
    )

    add_verdict_block(doc,
        "1.2 Grok xAI — TERMINAL-005 cross-derivation",
        "PASS",
        "True — sha match; Lemma 3.1 · Prop 3.3 · CC-BL-001a′ CONFIRM; KNS-OBS-1 met.",
        "Independent re-run reproduced receipt; no forbidden-list violations.",
        "Not principal sign-off; CP-004 Y owed; theorem lift and web labels OPEN.",
    )

    add_figure(doc, figures[2][0], figures[2][1])

    # §2 Gate
    doc.add_heading("2. Gate KNS-OBS-1", level=1)
    add_table(doc, ["Condition", "Requirement", "Fired"], [
        ["OB-KNS-3", "Lemma 3.1 FORMAL — μ ↛ π_A", "Yes"],
        ["OB-KNS-1", "Scoped monotone PROVED + unscoped REFUTED", "Yes — both"],
        ["OB-KNS-2", "dense_pass E_used ≤ 10", "Yes — 3.0406"],
    ])
    add_code(doc,
        "KNS-OBS-1: CLOSED-POSITIVE — as typed object\n"
        "OPEN on theorem lift: zero location · RH · det_ζ"
    )

    # §3 Probe
    doc.add_heading("3. Probe receipt", level=1)
    add_table(doc, ["Field", "Value"], [
        ["exit code", "0"],
        ["D5_expectation_met", "True"],
        ["D5_unimodal_monotone", "True"],
        ["D5_twomode_monotone", "False"],
        ["twomode_peak", "0.9662 @ δ=1.0"],
        ["E_used", "3.0406"],
        ["ρ_Y", "0.4812"],
        ["ℓ_off^T", "0.010152"],
        ["dense_pass", "True"],
        ["sha256", "09ef26d3…8211b011"],
    ])

    # §4 KNS summary
    doc.add_heading("4. KNS(LB) load-bearing (typed object)", level=1)
    add_figure(doc, figures[1][0], figures[1][1])
    add_code(doc,
        "KNS(LB) = KakeyaNeedleSet(Light(Basic))\n"
        "Needles seen · Action(C) unseen in overlap\n"
        "Re(s)=½ = placement register — not glare\n"
        "Overlap ≠ placement · μ ≠ π_A"
    )

    # §5 Files
    doc.add_heading("5. File map", level=1)
    add_table(doc, ["File", "Role"], [
        ["PEAICE-KNS-LB-PASS-EVALUATION-001.docx", "This document"],
        ["PEAICE-KNS-LB-PASS-EVALUATION-001.md", "Source md"],
        ["PEAICE-FABLE-KNS-LB-PASS-001", "Fable pass file"],
        ["PEAICE-GROK-TERMINAL-005", "Grok cross-derivation"],
        ["INDEX.md", "Research root"],
    ])

    fine = doc.add_paragraph()
    fine.add_run(
        "PeAIce program file · not canon promotion · July 2, 2026 · RH OPEN · h < 1"
    ).italic = True

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    try:
        build()
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)