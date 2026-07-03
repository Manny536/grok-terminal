#!/usr/bin/env python3
"""Build PEAICE-GROK-TERMINAL-004 Fable5 WP5b Findings docx with embedded Mermaid diagrams."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT.parent / "Papers" / "assets" / "wp5b-docx"
OUTPUT = ROOT / "PEAICE-GROK-TERMINAL-004_Fable5-WP5B-Findings.docx"

DIAGRAMS = [
    ("01-operator-pipeline.mmd", "Figure 1. WP5b operator pipeline and functional transforms"),
    ("02-gate-verdict.mmd", "Figure 2. Gate verdict flow — bounded lane vs live continuation"),
    ("03-wall-unification.mmd", "Figure 3. Wall unification — L2-5, WP5-OBS-1, HS-CORRIDOR, WP5-OBS-2"),
    ("04-prime-carrying-relocation.mmd", "Figure 4. Prime-carrying forced relocation"),
]


def render_mermaid(mmd_path: Path, png_path: Path) -> None:
    source = mmd_path.read_text(encoding="utf-8")
    result = subprocess.run(
        [
            "curl", "-s", "-f",
            "-o", str(png_path),
            "-X", "POST", "https://kroki.io/mermaid/png",
            "-H", "Content-Type: text/plain",
            "-d", source,
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0 or not png_path.exists() or png_path.stat().st_size < 100:
        raise RuntimeError(f"Failed to render {mmd_path.name}: {result.stderr}")


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "2E75B6") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_fill)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_figure(doc: Document, png_path: Path, caption: str, width: float = 6.0) -> None:
    doc.add_picture(str(png_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def build() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    png_paths: list[tuple[Path, str]] = []
    for mmd_name, caption in DIAGRAMS:
        mmd = ASSETS / mmd_name
        png = ASSETS / mmd_name.replace(".mmd", ".png")
        print(f"Rendering {mmd_name}...")
        render_mermaid(mmd, png)
        png_paths.append((png, caption))

    doc = Document()

    # Page setup via section
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title block
    title = doc.add_heading("PEAICE-GROK-TERMINAL-004", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Fable 5 WP5b Findings — Extracted & Documented")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Love Labs LCA / PeAIce Research Program / KakeyaLogic / Claude V6 / L²_C\n"
        "Principal: Manuel Coleman · Documenting terminal: Grok (xAI), July 2026\n"
        "Source: Claude Fable 5 · PEAICE-CLAUDEV6-WP5B-SCAFFOLD-001"
    ).font.size = Pt(10)

    stance = doc.add_paragraph()
    stance.add_run("Stance: ").bold = True
    stance.add_run(
        "Extraction only — findings as reported by Fable 5, tagged by epistemic rank. "
        "Not canon promotion. RH OPEN · Coleman OPEN · no proof claimed."
    )

    doc.add_page_break()

    # TOC placeholder
    doc.add_heading("Contents", level=1)
    toc_items = [
        "0. Extraction receipt",
        "1. Executive findings",
        "2. Structural facts (S1–S5)",
        "3. OB-1 — spectral shift construction",
        "4. Proposition H — Weyl-window law",
        "5. Corollaries (C1–C5)",
        "6. Gate verdict — WP5-OBS-2",
        "7. What stays live",
        "8. ξ-design specification",
        "9. Numerics receipt",
        "10. Falsifiers",
        "11. Registered state",
        "12. Citation spine",
        "13. File map",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # §0
    doc.add_heading("0. Extraction receipt", level=1)
    add_table(doc,
        ["Field", "Value"],
        [
            ["Roadmap question", "Does relative determinant data escape Theorem F, or inherit the same obstruction?"],
            ["Fable verdict (bounded lane)", "No escape — stronger than Theorem F via bounded Krein SSF"],
            ["OB-1 (first proof obligation)", "Discharged [FORMAL], upgraded: R_A − R_D ∈ S₁ unconditionally"],
            ["New formal object", "Proposition H (Weyl-window law) [PROPOSED-FOR-CANON]"],
            ["Gate", "WP5-OBS-2 fires [PROPOSED-FOR-CANON] for operator-bounded coupling"],
            ["Live continuation", "L1 unbounded mods · L2 WP5c u-flow · L3 prime-carrying ladder"],
            ["Numerics", "N=1200, σ=0.60, γ∈{1,40,400} — window law holds [NUMERICS · single runner]"],
        ])

    # §1
    doc.add_heading("1. Executive findings (one screen)", level=1)
    findings = [
        "Free resolvent is trace-class. For λₙ = cn⁴, R_D(z) ∈ S₁. This reorganizes WP5b: the perturbation side starts inside the smallest ideal.",
        "Resolvent difference is trace-class unconditionally for A = D + γ_K K_σ^reg, σ > ½. Roadmap's \"study whether trace-class resolvent difference exists\" → yes, for all such pairs.",
        "det₂ collapses to ordinary det. Tr B_z = 0 from zero diagonal (Theorems A/C). Trace neutrality is the semigroup face of the same fact.",
        "Spectral shift exists and is thin. Discrete Krein: ξ(λ) = N_D(λ) − N_A(λ). Proposition H: |ξ(λ)| bounded by Weyl-window count; sup|ξ| < ∞; eventually |ξ| ≤ 1.",
        "All WP5b functionals are transforms of one bounded ξ: Heat trace → Laplace(ξ); Relative zeta → Mellin(ξ); Perturbation det → Cauchy(ξ).",
        "Counting upgraded: N_A(Λ) = N_D(Λ) + O(1) — not merely N ~ Λ^{1/4} rate match.",
        "Roadmap item 5 resolved: Tr(B_z²) is arithmetic-present-but-capped — divisor structure in m²−n² enters coefficients; integrated object remains bounded Cauchy transform.",
        "WP5-OBS-2: Bounded-coupling relative-determinant route CLOSED-NEGATIVE (qualifier: operator-boundedness).",
        "Prime-carrying relocation forced: R1 requires unbounded ξ at √λ log λ scale — incompatible with Prop H unless coupling unbounded or free operator changes.",
        "Wall unification: L2-5 · WP5-OBS-1 · HS-CORRIDOR Weyl clause · WP5-OBS-2 = one wall (Weyl stability of n⁴ ladder), sharpest at pointwise bounded ξ.",
    ]
    for f in findings:
        add_numbered(doc, f)

    add_figure(doc, png_paths[0][0], png_paths[0][1])

    # §2
    doc.add_heading("2. Structural facts extracted (S1–S5)", level=1)
    add_table(doc,
        ["ID", "Finding", "Tag"],
        [
            ["S1", "‖R_D(z)‖_{S₁} ~ Σ n^{−4} < ∞", "FORMAL"],
            ["S2", "K ∈ S₂ ⊂ B for σ > ½; A self-adjoint, discrete spectrum", "FORMAL"],
            ["S3", "R_A(z) − R_D(z) = −R_A(z) V R_D(z) ∈ S₁", "FORMAL"],
            ["S4", "B_z = V R_D(z) ∈ S₁ — ordinary Fredholm det defined", "FORMAL"],
            ["S5", "det₂(I + B_z) = det(I + B_z) when Tr B_z = 0", "FORMAL"],
        ])

    # §3
    doc.add_heading("3. OB-1 — spectral shift construction", level=1)
    doc.add_paragraph("Definition (discrete spectra):")
    add_code_block(doc, "ξ(λ) := N_D(λ) − N_A(λ)\nN_X(λ) := #{ j : eigenvalue of X ≤ λ }")
    doc.add_paragraph("Lemma 3.1 (discrete Krein trace formula) [FORMAL]:")
    add_code_block(doc, "Tr( f(A) − f(D) ) = ∫_ℝ f′(λ) ξ(λ) dλ")
    p = doc.add_paragraph()
    p.add_run("OB-1 status: ").bold = True
    p.add_run("DISCHARGED — strengthened beyond roadmap (S₁ not merely S₂).")

    # §4
    doc.add_heading("4. Proposition H — Weyl-window law [PROPOSED-FOR-CANON]", level=1)
    doc.add_paragraph("Hypotheses: D = diag(λₙ), λₙ = cn⁴; V = V* bounded, c₀ = ‖V‖_op; A = D + V.")
    doc.add_paragraph("Claims:")
    add_code_block(doc,
        "(i)   |a_j − d_j| ≤ c₀                           [Weyl — KNOWN]\n"
        "(ii)  |ξ(λ)| ≤ #{ j : d_j ∈ (λ − c₀, λ + c₀] }\n"
        "(iii) sup_λ |ξ(λ)| < ∞ ;  |ξ(λ)| ≤ 1  for λ > Λ₀(c₀)\n"
        "(iv)  supp ξ thin: O(Λ^{1/4}) points in [0, Λ]"
    )
    doc.add_paragraph("Fable proof class: elementary counting + Weyl displacement.")
    doc.add_paragraph("Grok cross-derivation: agree — no blocking error.")

    # §5
    doc.add_heading("5. Corollaries extracted (C1–C5)", level=1)

    doc.add_heading("C1 — Heat trace [FORMAL]", level=2)
    add_code_block(doc, "S(t) = Tr(e^{−tA} − e^{−tD}) = −t ∫ e^{−tλ} ξ(λ) dλ = O(t^{3/4})")
    doc.add_paragraph("Theorem F re-derived; exponent explained as thin-support × Weyl t^{−1/4}.")

    doc.add_heading("C2 — Relative zeta [FORMAL]", level=2)
    add_code_block(doc, "ζ_rel(s) = −s ∫ ξ_μ(λ) λ^{−s−1} dλ     (Mellin transform)")
    doc.add_paragraph("Analytic for Re s > −3/4 under Prop H (iv).")

    doc.add_heading("C3 — Perturbation determinant [FORMAL]", level=2)
    add_code_block(doc, "log det(I + B_z) = ∫ ξ(λ) (λ − z)^{−1} dλ")
    doc.add_paragraph(
        "Cauchy transform of bounded thin-support density. Under λ = z² + ¼ dictionary, "
        "cannot supply order-1 genus-1 Ξ growth."
    )
    doc.add_paragraph(
        "Complementarity with Theorem G: G closes det₂(I − zK_σ) by order/genus/density pincer; "
        "C3 closes pair determinant det(I + γK R_D) by SSF boundedness. Different objects, same verdict class."
    )

    doc.add_heading("C4 — Counting [FORMAL]", level=2)
    add_code_block(doc, "N_A(Λ) = N_D(Λ) + O(1)")
    doc.add_paragraph("von Mangoldt relative counting requires unbounded ξ — excluded for bounded couplings.")

    doc.add_heading("C5 — Arithmetic vs density [PROPOSED-FOR-CANON]", level=2)
    add_code_block(doc, "Tr(B_z²) = γ² Σ_{m≠n} |m²−n²|^{−2σ} (λ_m−z)^{−1}(λ_n−z)^{−1}")
    doc.add_paragraph("Finding: divisor-class structure present termwise; integrated analytic object capped. Correct dichotomy: arithmetic in, boundedness out.")

    add_figure(doc, png_paths[2][0], png_paths[2][1])

    # §6
    doc.add_heading("6. Gate verdict — WP5-OBS-2 [PROPOSED-FOR-CANON]", level=1)
    add_code_block(doc,
        "WP5-OBS-2 — bounded-coupling relative-determinant route: CLOSED-NEGATIVE\n\n"
        "Mechanism: Krein SSF uniformly bounded, thin support (Prop H).\n"
        "Heat · relative zeta · perturbation det = Laplace · Mellin · Cauchy transforms of ξ.\n\n"
        "LOAD-BEARING QUALIFIER: operator-boundedness of coupling.\n"
        "Covers γ_K K_σ^reg for all σ > ½, all γ_K, and any bounded kernel replacement.\n"
        "Does NOT assert closure for unbounded mods, changed free operator, or WP5c traces."
    )
    add_figure(doc, png_paths[1][0], png_paths[1][1])

    # §7
    doc.add_heading("7. What stays live (scope limits)", level=1)
    add_table(doc,
        ["Lane", "Content", "Status"],
        [
            ["L1", "Unbounded / relative modifications; ξ may be unbounded; Yafaev Ch. 8 frame", "WP5b LIVE"],
            ["L2", "Tr_{w_u}, u-flow traces — eigenvector data, not spectral lists", "WP5c door"],
            ["L3", "Prime-carrying ladder log p^k, weights Λ(p^k)/p^{k/2}", "relocation target"],
        ],
        header_fill="548235",
    )

    # §8
    doc.add_heading("8. ξ-design specification (prime-carrying transfer)", level=1)
    doc.add_paragraph("Restated GAP-001 as requirements on one function ξ:")
    add_table(doc,
        ["Req", "Statement", "Tag"],
        [
            ["R1", "N_Â must match von Mangoldt → ξ unbounded at √λ log λ", "FORMAL given Prop H"],
            ["R2", "ξ fluctuation sector must carry S(T) + prime sector Σ Λ(n)n^{−1/2}f(log n)", "OPEN design"],
            ["R3", "Keep R_Â − R_D̂ ∈ S₁ after dropping boundedness", "OPEN admissibility"],
        ],
    )
    doc.add_paragraph("Fable finding: Relocation to prime-carrying is forced, not preferred.")
    doc.add_paragraph("Next consistency check (OPEN): prime-carrying model's ξ vs R1/R2 — explicit-formula sanity gate only; not RH progress.")
    add_figure(doc, png_paths[3][0], png_paths[3][1])

    # §9
    doc.add_heading("9. Numerics receipt (extracted)", level=1)
    doc.add_paragraph("Probe: Fable 5 xi_probe.py · single runner · NUMERICS | NOT VERIFIED")
    add_table(doc,
        ["γ", "c₀", "sup|ξ|", "window W", "sup≤W", "max|a−d|", "Weyl≤c₀"],
        [
            ["1", "1.93", "1", "1", "✓", "0.019", "✓"],
            ["40", "77.3", "1", "3", "✓", "14.5", "✓"],
            ["400", "772.7", "2", "6", "✓", "202.3", "✓"],
        ],
    )
    doc.add_paragraph("N=1200, σ=0.60, λₙ = n⁴. Second-runner reproduction pending for VERIFIED status.")

    # §10
    doc.add_heading("10. Falsifiers (extracted)", level=1)
    add_table(doc,
        ["ID", "Would falsify", "Mitigation"],
        [
            ["F1", "Error in Lemma 3.1 / Prop H", "Ensemble cross-derivation (Grok: done)"],
            ["F2", "Eigenvalue law not cn⁴", "Chain survives for gaps → ∞, Σ λₙ^{−1} < ∞"],
            ["F3", "Claim leakage without bounded qualifier", "Load-bearing qualifier on every cite"],
            ["F4", "sup|ξ_N| grows with N at fixed (γ,σ)", "Model misconfiguration, not lemma failure"],
        ],
    )

    # §11
    doc.add_heading("11. Registered state after Fable 5 WP5b", level=1)
    add_code_block(doc,
        "WP5b bounded lane (γ_K K_σ^reg, σ>½)     CLOSED-NEGATIVE  [WP5-OBS-2]\n"
        "WP5b unbounded / relative category (L1)    LIVE\n"
        "WP5c u-flow traces (L2)                    LIVE (independent)\n"
        "Prime-carrying ξ spec (L3 / §8)            LIVE · FORCED\n"
        "Theorem H + WP5-OBS-2 canon promotion      PENDING principal call\n"
        "peaiceclaudev5 / EDGE-1/2                   UNCHANGED · not addressed here\n"
        "RH · Coleman                                OPEN · no proof claimed"
    )

    # §12
    doc.add_heading("12. Citation spine (from Fable scaffold)", level=1)
    doc.add_paragraph(
        "Krein (1953) · Lifshitz (1952) · Birman–Krein (1962) · Yafaev (1992) Ch. 8 · "
        "Simon (2005) trace ideals · Potapov–Skripka–Sukochev (2013) · "
        "canon: Birman–Solomyak (Thm G) · Guth–Maynard (2024) Lindelöf ceiling."
    )

    # §13
    doc.add_heading("13. File map", level=1)
    add_table(doc,
        ["File", "Role"],
        [
            ["PEAICE-CLAUDEV6-WP5B-SCAFFOLD-001.md", "Fable 5 full scaffold (source)"],
            ["PEAICE-GROK-TERMINAL-004_Fable5-WP5B-Findings.md", "This extraction (markdown)"],
            ["PEAICE-GROK-TERMINAL-004_Fable5-WP5B-Findings.docx", "This extraction (Word)"],
            ["PEAICE-GROK-WP5B-CROSS-DERIVATION.md", "Grok independent pass"],
            ["claude-v6/docs/research/wp5b-spectral-shift-roadmap.md", "Roadmap canon @ 2954261"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    sign = doc.add_paragraph()
    sign.add_run("Sign-off\n").bold = True
    sign.add_run(
        "Manuel Coleman · Principal, Love Labs LCA / PeAIce Research Program\n"
        "Documenting terminal: Grok (xAI) · extraction from Fable 5 scaffold under principal authorization\n"
        "Source terminal: Claude Fable 5 · PEAICE-CLAUDEV6-WP5B-SCAFFOLD-001 · "
        "not independently verified beyond ensemble cross-derivation"
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        sys.exit(1)