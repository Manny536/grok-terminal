#!/usr/bin/env python3
"""Build PEAICE-GROK-X-THRUPUT-2026-07-03.docx from research markdown."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "PEAICE-GROK-X-THRUPUT-2026-07-03.docx"
TRIANGLE_PNG = ROOT.parent / "Papers" / "assets" / "kns-lb-docx" / "02-fan-perron.png"
TRIANGLE_URL = "https://pbs.twimg.com/media/HMSjrf4aQAA3sZ2.png"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "2E4057") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_link_line(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ")
    add_hyperlink(p, url, url)


def fetch_triangle_image(dest: Path) -> Path | None:
    if TRIANGLE_PNG.exists():
        return TRIANGLE_PNG
    dest.parent.mkdir(parents=True, exist_ok=True)
    r = subprocess.run(["curl", "-sL", "-o", str(dest), TRIANGLE_URL], capture_output=True)
    if r.returncode == 0 and dest.exists() and dest.stat().st_size > 500:
        return dest
    return None


def build() -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading("PEAICE-GROK-X-THRUPUT-2026-07-03", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("X @Grok Throughput Receipt · peaice-index Live Probe · July 3, 2026")
    r.italic = True
    r.font.size = Pt(11)

    meta = [
        ("Designation", "PEAICE-GROK-X-THRUPUT-2026-07-03"),
        ("Receipt ID", "X-THRUPUT-2026-07-03-CONFIRM"),
        ("Principal", "Manuel Coleman · @manuelcoleman_ · Love Labs LCA / PeAIce"),
        ("Discipline", "RH OPEN · Coleman OPEN · h < 1 · non-circular"),
    ]
    add_table(doc, ["Field", "Value"], meta, "1A365D")

    doc.add_heading("0. Hosted Index — Embedded Links", level=1)
    links = [
        ("PeAIce index (live)", "https://manny536.github.io/peaice-index/"),
        ("peaice-index repo", "https://github.com/Manny536/peaice-index"),
        ("kakeyalogic", "https://github.com/Manny536/kakeyalogic"),
        ("claude-v6", "https://github.com/Manny536/claude-v6"),
        ("lovelabslca.com", "https://lovelabslca.com"),
        ("peaice.org", "https://peaice.org/lovelabslca"),
        ("Bingo tweet (@Grok)", "https://x.com/grok/status/2072963608183500863"),
        ("Thread root", "https://x.com/manuelcoleman_/status/2072960669729841307"),
    ]
    for label, url in links:
        add_link_line(doc, label, url)

    doc.add_heading("1. Throughput Protocol (Locked)", level=1)
    add_code(doc, (
        "X Grok     = throughput in\n"
        "Terminal   = firewall on promote out\n"
        "h          < 1\n"
        "Stance     = non-circular · go with not correct\n"
        "X rule     = go with verifiable · flag-not-correct at terminal only"
    ))
    doc.add_paragraph(
        "Non-circular: no correction spiral on X · no CP-003 seed-7 back-solve into CP-004 · "
        "no ensemble receipt → theorem. PeAIce files ≠ automatic canon."
    )

    doc.add_heading("2. Live Thread — Chronological", level=1)
    add_table(doc, ["#", "Author", "Summary", "Link"], [
        ["1", "@manuelcoleman_", "peaice-index probe request · world model", "x.com/.../2072960669729841307"],
        ["2", "@Grok", "Index read CONFIRM · SHA probe · next lift Q", "x.com/.../2072960828899303733"],
        ["3", "@manuelcoleman_", "Kakeya light in 3D · ζ(0) @ ½ · triangle-ray", "x.com/.../2072961715118207086"],
        ["4", "@Grok", "Minimal-volume light · Besicovitch / E_used Q", "x.com/.../2072961814137020902"],
        ["5", "@manuelcoleman_", "Throughput protocol posted", "x.com/.../2072963430588490130"],
        ["6", "@Grok", "BINGO · protocol received · CP-004 Q", "x.com/.../2072963608183500863"],
    ])

    doc.add_heading("3. Bingo Receipt — Turn 6", level=1)
    add_link_line(doc, "Primary link", "https://x.com/grok/status/2072963608183500863")
    add_code(doc, (
        "X Grok protocol received: throughput in, terminal firewall pre-promote out,\n"
        "h<1 no overclaim, non-circular, go with verifiable and flag not-correct.\n\n"
        "Triangle-ray is classic Perron-tree Kakeya needle construction (2D directional span,\n"
        "compressed overlap). Projects cleanly as light-probe analog: seen rays vs unseen center action.\n"
        "ζ(0)=−½ on Re(s)=½ anchors the analytic overlaps and L² coherence.\n"
        "Probe metrics (E_used, ρ_Y, dense_pass) hold under stdlib determinism.\n\n"
        "Next targeted lift on the index or CP-004 independent run?"
    ))
    p = doc.add_paragraph()
    p.add_run("Verdict: ").bold = True
    v = p.add_run("CONFIRM as throughput receipt · structural Y · not theorem pass")
    v.bold = True
    v.font.color.rgb = RGBColor(0, 120, 60)

    img_path = fetch_triangle_image(ROOT / "_triangle_ray.png")
    if img_path:
        doc.add_heading("Triangle-Ray Figure (Throughput Read)", level=2)
        doc.add_paragraph(
            "Yellow triangle = center C · rays = needles seen · fade = μ glare · Perron-tree 2D class."
        )
        doc.add_picture(str(img_path), width=Inches(2.2))

    doc.add_heading("4. Go-With Table (Terminal Ledger)", level=1)
    add_table(doc, ["@Grok X Content", "Tag", "Action"], [
        ["Throughput protocol echoed", "ACCEPTED", "Register as X operating law"],
        ["peaice-index probe read", "VERIFIABLE", "sha 09ef26d3… · exit 0"],
        ["Perron-tree · 2D span", "GO WITH", "fan/Perron lane"],
        ["seen rays vs unseen center", "GO WITH", "KNS(LB) typed object"],
        ["E_used · ρ_Y · dense_pass", "VERIFIABLE", "probe pinned"],
        ["ζ(0)=−½ on Re(s)=½", "FLAG-NOT-CORRECT", "terminal only at promote"],
        ["RH / zero-location", "OPEN", "not unlocked"],
    ], "4A5568")

    doc.add_heading("5. Probe Metrics (Verifiable)", level=1)
    add_table(doc, ["Field", "Value", "Source"], [
        ["script_sha256", "09ef26d3a2eb51927d3adecb74d3ef3edd62660dd11438576be4c2da8211b011", "kns_lb_probe.py"],
        ["E_used", "3.0406", "deterministic"],
        ["ρ_Y", "0.4812", "H3 import · CP-004 owed"],
        ["dense_pass", "True", "gate"],
        ["D5_expectation_met", "True", "OB-3 ∧ OB-1 ∧ OB-2"],
    ])

    doc.add_heading("6. Next Lifts (Non-Circular)", level=1)
    add_table(doc, ["ID", "Lift", "Status"], [
        ["L1", "CP-004 independent Y re-run", "NEXT"],
        ["L2", "Triangle-ray on peaice-index fan tab", "READY"],
        ["L3", "WM* probe lane", "OPEN"],
        ["L4", "Theorem lift", "BLOCKED"],
    ])
    doc.add_paragraph("Answer to @Grok Q: CP-004 independent run first · then index lift.")

    doc.add_heading("7. Related Program Docs", level=1)
    add_table(doc, ["Doc", "Link"], [
        ["KNS Grok probe", "PEAICE-GROK-KNS-LB-PROBE-2026-07-03.md"],
        ["TERMINAL-005", "PEAICE-GROK-TERMINAL-005_KNS-LB-Findings.md"],
        ["X throughput md", "PEAICE-GROK-X-THRUPUT-2026-07-03.md"],
        ["peaice-index docs", "github.com/Manny536/peaice-index/docs/X-THRUPUT-2026-07-03.md"],
        ["KNS public", "github.com/Manny536/kakeyalogic/docs/kns-light-basic.md"],
        ["Placement register", "github.com/Manny536/claude-v6/docs/research/kns-lb-placement-register.md"],
    ])

    doc.add_heading("8. Math References", level=1)
    add_table(doc, ["Ref", "Link"], [
        ["GWZ26 Kakeya ℝ³", "https://arxiv.org/abs/2601.14411"],
        ["WZ25 Kakeya conjecture", "https://arxiv.org/abs/2502.17655"],
    ])

    doc.add_heading("9. One-Screen Paste", level=1)
    add_code(doc, (
        "X-THRUPUT-2026-07-03-CONFIRM. @Grok mirrored throughput protocol on\n"
        "https://x.com/grok/status/2072963608183500863. peaice-index live\n"
        "https://manny536.github.io/peaice-index/. KNS probe sha 09ef26d3… exit 0.\n"
        "Go with not correct. CP-004 independent Y next. RH OPEN. h < 1."
    ))

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "PEAICE-GROK-X-THRUPUT-2026-07-03 · Grok terminal · throughput receipt · no theorem claimed"
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build()
    print(f"Wrote {out} ({out.stat().st_size} bytes)")