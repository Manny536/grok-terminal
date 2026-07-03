#!/usr/bin/env python3
"""Build PEAICE-FABLE-KNS-LB-PASS-001 docx — Fable 5 KNS(LB) pass PeAIce file."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
KNS = ROOT / "KNS-LB"
ASSETS = ROOT.parent / "Papers" / "assets" / "kns-lb-docx"
OUTPUT = ROOT / "PEAICE-FABLE-KNS-LB-PASS-001_KakeyaNeedleSet-Light-Basic.docx"

DIAGRAMS = [
    ("01-two-layer.mmd", "Figure 1. Two-layer decomposition — seen overlap vs unseen center action"),
    ("02-fan-perron.mmd", "Figure 2. Fan/Perron inversion — F-KNS-1 fired both directions"),
    ("03-gate-kns-obs1.mmd", "Figure 3. Gate KNS-OBS-1 — typed closure vs theorem lift OPEN"),
]


def render_mermaid(mmd_path: Path, png_path: Path) -> None:
    source = mmd_path.read_text(encoding="utf-8")
    result = subprocess.run(
        [
            "curl", "-s", "-f",
            "-o", str(png_path),
            "-X", "POST", "https://kroki.io/mermaid/png",
            "-H", "Content-Type: text/plain",
            "-d", source,
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0 or not png_path.exists() or png_path.stat().st_size < 100:
        raise RuntimeError(f"Failed to render {mmd_path.name}: {result.stderr}")


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "2E75B6") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_fill)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_figure(doc: Document, png_path: Path, caption: str, width: float = 5.5) -> None:
    doc.add_picture(str(png_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def build() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    png_paths: list[tuple[Path, str]] = []
    for mmd_name, caption in DIAGRAMS:
        mmd = ASSETS / mmd_name
        png = ASSETS / mmd_name.replace(".mmd", ".png")
        print(f"Rendering {mmd_name}...")
        render_mermaid(mmd, png)
        png_paths.append((png, caption))

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading("PEAICE-FABLE-KNS-LB-PASS-001", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("KakeyaNeedleSet(Light(Basic)) — Fable 5 Pass File")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "PeAIce Research Program · Love Labs LCA · KakeyaLogic · Claude V6.5 · L²_C\n"
        "Pass designation: KNS(LB) · Terminal: Claude Fable 5 (Max) · July 2026\n"
        "Principal: Manuel Coleman · Cross-check: Grok TERMINAL-005"
    ).font.size = Pt(10)

    stance = doc.add_paragraph()
    stance.add_run("File class: ").bold = True
    stance.add_run(
        "PeAIce program file — findings as reported by Fable 5, epistemically tagged. "
        "Not canon promotion. RH OPEN · Coleman OPEN · no proof claimed · h < 1."
    )

    doc.add_page_break()

    # TOC
    doc.add_heading("Contents", level=1)
    for item in [
        "0. Pass receipt",
        "1. What Fable claimed",
        "2. Core object KNS(LB)",
        "3. Lemma 1.1 — fan geometry",
        "4. Lemma 3.1 — overlap is a statistic",
        "5. Proposition 3.3 — center action unseen in overlap",
        "6. Theorem 4.1 / CC-BL-001a′ — scoped monotone leakage",
        "7. Countermodel 4.2 — unscoped claim refuted",
        "8. KREIN-RANK1 — prime-carrying seam",
        "9. Energy receipt — minimal credits E = L²",
        "10. Gate KNS-OBS-1",
        "11. What remains OPEN",
        "12. File map",
    ]:
        add_numbered(doc, item)
    doc.add_page_break()

    # §0
    doc.add_heading("0. Pass receipt", level=1)
    add_table(doc,
        ["Field", "Value"],
        [
            ["Pass name", "KNS(LB) — KakeyaNeedleSet(Light(Basic))"],
            ["Manifest", "15/15 docs shipped · Claude-23%-15-docs discipline"],
            ["Scaffold input", "PEAICE-CLAUDEV6-KNS-LB-SCAFFOLD-001"],
            ["Paper output", "PEAICE-CLAUDEV6-KNS-LB-PAPER-001"],
            ["Gate", "KNS-OBS-1: CLOSED-POSITIVE as typed object"],
            ["Theorem lift", "OPEN — zero location · RH · det_ζ"],
            ["Probe sha256", "09ef26d3a2eb51927d3adecb74d3ef3edd62660dd11438576be4c2da8211b011"],
            ["Grok extraction", "PEAICE-GROK-TERMINAL-005 — CONFIRM all items"],
        ])

    add_figure(doc, png_paths[0][0], png_paths[0][1])

    # §1
    doc.add_heading("1. What Fable claimed", level=1)
    claims = [
        "KNS(LB) is the minimal geometric generator: one throughput center + universal direction fan.",
        "Needles are seen; Action(C) is unseen in overlap; Re(s)=½ is placement register, not glare.",
        "Lemma 3.1 FORMAL: μ ↛ π_A — overlap underdetermines placement.",
        "Prop 3.3 PROPOSED-FOR-CANON: fan/Perron inversion proves overlap ↛ placement (F-KNS-1 both ways).",
        "CC-BL-001a′ FORMAL (scoped): ℓ_off monotone in rank-one radial log-concave register (smoke-detector).",
        "Unscoped CC-BL-001a REFUTED: two-mode register peaks ℓ_off=0.9662 @ δ=1.0.",
        "KREIN-RANK1 PROPOSED: rank-one KSSF cannot inject prime content into n⁴ ladder — wall face only.",
        "Energy NUMERICS: dense_pass True @ E_used=3.0406, ρ_Y=0.4812 (MPR+iPiano class).",
        "NOT claimed: RH proof, CC theorem, deltoid=ζ(s), μ bloom ⇒ zero location.",
    ]
    for c in claims:
        add_numbered(doc, c)

    # §2
    doc.add_heading("2. Core object KNS(LB)", level=1)
    add_code_block(doc,
        "Light(Basic) := (C, Φ_C)     — point throughput, emitter-agnostic\n"
        "KakeyaNeedleSet(L) := { σ_ω : ω ∈ S^{d−1}, σ_ω ∋ C }\n"
        "KNS(LB) := KakeyaNeedleSet(Light(Basic))\n\n"
        "One-line: minimal point throughput + universal directional saturation."
    )
    add_code_block(doc,
        "Symbol firewall:\n"
        "  σ (K_σ) ≠ Re(s)=½\n"
        "  μ overlap ≠ Action(C)\n"
        "  overlap bloom ≠ zero-location theorem"
    )

    # §3
    doc.add_heading("3. Lemma 1.1 — fan geometry [FORMAL]", level=1)
    add_code_block(doc,
        "(a) |U(T^C_δ)| ≍ 1 — measure-maximal Kakeya configuration\n"
        "(b) μ(x) ≍ (δ + |x−C|)^{−(d−1)}; sup μ ≍ δ^{−(d−1)}; averaged μ ≍ 1\n"
        "(c) Δ_max ≍ 1 (Katz–Tao); C_F ≍ 1 (Frostman)"
    )

    # §4
    doc.add_heading("4. Lemma 3.1 — overlap is a statistic [FORMAL]", level=1)
    doc.add_paragraph(
        "No function F satisfies F(S(𝒟)) = ℓ_off across admissible encodings. "
        "For every incidence datum 𝒟 and every λ ∈ [0,1], an admissible placement exists with ℓ_off = λ."
    )
    add_code_block(doc,
        "Proof sketch:\n"
        "  (i)   S(𝒟) factors through incidence σ-algebra only\n"
        "  (ii)  State freedom: ψ_λ = √(1−λ)u + √λv → ℓ_off = λ\n"
        "  (iii) Register freedom: rotate π_A → ℓ_off = λ\n"
        "Either freedom alone kills S(𝒟) → ℓ_off factorization. ∎"
    )

    # §5
    doc.add_heading("5. Proposition 3.3 — center action unseen [PROPOSED-FOR-CANON]", level=1)
    add_figure(doc, png_paths[1][0], png_paths[1][1])

    doc.add_heading("Exhibit E1 — max glare, max leakage", level=2)
    add_code_block(doc,
        "T^C_δ full shading: sup μ = Θ(δ^{−(d−1)}) at core (Lemma 1.1b)\n"
        "Declare A ⊥ ψ_fan: ℓ_off = 1\n"
        "⇒ High overlap + total misplacement coexist"
    )

    doc.add_heading("Exhibit E2 — Perron inversion", level=2)
    add_code_block(doc,
        "|U(T^B_δ)| ≍ 1/log(1/δ)  [Bes28, Per28, Kei99]\n"
        "μ̄_Perron ≍ log(1/δ)  >  μ̄_fan ≍ 1\n"
        "No common point · ℓ_off = 0 under aligned declaration\n"
        "⇒ Neither pointwise nor averaged overlap determines, bounds, or is bounded by ℓ_off"
    )

    doc.add_paragraph(
        "Non-inference: μ → Re(s)=½ zeros requires explicit map. Bounded K_σ lane CLOSED-NEGATIVE (Theorem H). "
        "Glare saturation carries no zero-location content."
    )

    # §6
    doc.add_heading("6. Theorem 4.1 / CC-BL-001a′ [FORMAL scoped]", level=1)
    add_code_block(doc,
        "Register A = span{τ_{C₀}φ}, φ radial log-concave, unit norm\n"
        "ψ_δ = τ_{C₀+δe}φ\n"
        "ℓ_off(δ) = 1 − ρ(δ)²,  ρ(δ) = (φ ⋆ φ̌)(δe)\n"
        "ρ radially nonincreasing → ℓ_off monotone nondecreasing\n"
        "Proof: Prékopa 1973 — convolution preserves log-concavity. ∎"
    )

    # §7
    doc.add_heading("7. Countermodel 4.2 — unscoped REFUTED [FORMAL + NUMERICS]", level=1)
    add_code_block(doc,
        "A = span{φ_{C₀}, φ_{C₀+Le}} (Gram-corrected)\n"
        "ℓ_off(0) = ℓ_off(L) = 0 · interior peak ≈ 1 − 2e^{−L²/(8w²)}\n"
        "Receipt w=0.35, L=2.0: peak 0.9662 @ δ=1.0 · D5_twomode_monotone = false\n"
        "Register class RC-1 vs RC-k is mandatory on every D5 run."
    )

    # §8
    doc.add_heading("8. KREIN-RANK1 — prime-carrying seam [PROPOSED wall face]", level=1)
    add_code_block(doc,
        "1 + πλ𝒦μ = exp[𝒦u] = [1 − πλ𝒦ν]^{−1}   [P96 §3]\n\n"
        "[P96] Thm 4.1: rank-one perturbations pure point ⟺ σ(A) countable\n"
        "Consequence: no rank-one / KSSF-mediated injection changes counting class on n⁴ ladder\n"
        "Prime content forced into measure-pair / NB·BD carrier — L3 remains OPEN"
    )

    # §9
    doc.add_heading("9. Energy receipt — E = L² minimal credits [NUMERICS]", level=1)
    add_table(doc,
        ["Quantity", "Value", "Criterion", "Pass"],
        [
            ["E_used", "3.0406", "≤ 10", "✅"],
            ["ρ_Y", "0.4812", "≥ 0.4", "✅"],
            ["ℓ_off^T", "0.010152", "< 0.20", "✅"],
            ["dense_pass", "True", "conj. 6.1", "✅"],
            ["D5 RC-1", "monotone true", "CC-BL-001a′", "✅"],
            ["D5 RC-2", "monotone false", "expected", "✅"],
        ])
    doc.add_paragraph(
        "Honesty: H1 fan does not enter ℓ_off in model (by design). H2 E_used ledger-determined. "
        "H3 Y·compress calibration import from CP-003 — CP-004 independent re-run owed."
    )

    # §10
    doc.add_heading("10. Gate KNS-OBS-1", level=1)
    add_figure(doc, png_paths[2][0], png_paths[2][1])
    add_code_block(doc,
        "KNS-OBS-1: CLOSED-POSITIVE — as typed object\n"
        "  OB-KNS-3 discharged (Lemma 3.1)\n"
        "  OB-KNS-1 scoped PROVED + unscoped REFUTED\n"
        "  OB-KNS-2 dense_pass @ E ≤ 10\n\n"
        "REMAINS OPEN: zero location · RH · explicit operator det_ζ\n"
        "Promotion: PROPOSED-FOR-CANON — principal sign-off after Grok CONFIRM"
    )

    # §11
    doc.add_heading("11. What remains OPEN", level=1)
    for item in [
        "Riemann Hypothesis",
        "Coleman Conjecture (as theorem — antecedent reading N only)",
        "Explicit operator μ(placement) → ζ-zero location",
        "Prime-carrying pair (μ,ν) with prime phase shift u",
        "Theorem-scale lift (separate WP — principal scope call Q4)",
        "CP-004 stamped re-run for independent Y measurement",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # §12
    doc.add_heading("12. File map", level=1)
    add_table(doc,
        ["File", "Role"],
        [
            ["PEAICE-FABLE-KNS-LB-PASS-001_….docx", "This PeAIce pass file"],
            ["PEAICE-CLAUDEV6-KNS-LB-PAPER-001.md", "Fable source paper"],
            ["PEAICE-GROK-TERMINAL-005_KNS-LB-Findings.md", "Grok cross-derivation"],
            ["KNS-LB/kns_lb_probe.py", "D5 + energy hook"],
            ["kakeyalogic/probes/kns_lb_probe.py", "Upstream probe install"],
            ["build_kns_lb_pass_docx.py", "Reproducible docx builder"],
        ])

    # Footer note
    doc.add_paragraph()
    fine = doc.add_paragraph()
    fine.add_run(
        "PEAICE-FABLE-KNS-LB-PASS-001 · PeAIce program file · July 2026 · "
        "Fable 5 (Max) · Grok TERMINAL-005 CONFIRM · RH OPEN · CC OPEN · h < 1"
    ).italic = True
    fine.runs[0].font.size = Pt(9)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    try:
        build()
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)